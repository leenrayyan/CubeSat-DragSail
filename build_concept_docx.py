"""
Generate AnaMALY_AIDragSail_ConceptDocument.docx — corrected and renumbered.

Plain academic style: Calibri body, black headings, simple tables, no colored
banners. Sections numbered 1..12 (no leading 3.).
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = r"C:\Users\leenr\Desktop\AESH\AnaMALY_AIDragSail_ConceptDocument.docx"

# Image paths
BLOCK_DIAGRAM = (r"C:\Users\leenr\Downloads"
                 r"\AnaMalyhackathon-20260510T203316Z-3-001"
                 r"\AnaMalyhackathon\AnaMALY_AIDragSail_Phase1_BlockDiagram.png")
FIG_ALT = r"C:\Users\leenr\Desktop\AESH\results\altitude_vs_time.png"
FIG_SOLAR = r"C:\Users\leenr\Desktop\AESH\results\solar_forecast.png"
FIG_KPI = r"C:\Users\leenr\Desktop\AESH\results\kpi_comparison.png"
FIG_SCORE = r"C:\Users\leenr\Desktop\AESH\results\score_landscape.png"


def add_heading(doc, text, level=1):
    """Plain black heading, no theme color."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def add_numbered(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(it)
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def add_table(doc, header, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(10)
    # Data rows
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = "Calibri"
            r.font.size = Pt(10)
    if col_widths_cm:
        for ri in range(len(table.rows)):
            for ci, w in enumerate(col_widths_cm):
                table.rows[ri].cells[ci].width = Cm(w)
    return table


def add_pagebreak(doc):
    doc.add_page_break()


def add_figure(doc, image_path, caption, width_cm=15.0):
    """Insert a centered image followed by an italic caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if os.path.exists(image_path):
        run.add_picture(image_path, width=Cm(width_cm))
    else:
        r = p.add_run(f"[missing image: {image_path}]")
        r.italic = True
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(10)
    cr.font.name = "Calibri"


def add_assumptions_box(doc, text):
    """Render the 'Key Assumptions' callout as a single-cell shaded table."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Light Grid Accent 1"
    cell = t.rows[0].cells[0]
    cell.text = ""
    p1 = cell.paragraphs[0]
    r1 = p1.add_run("Key Assumptions")
    r1.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(10)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(text)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10)


def main():
    doc = Document()

    # ---- Document defaults ----
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0, 0, 0)

    # ---- Title block ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("AI-Driven Atmospheric Drag Sail for Rapid CubeSat Deorbiting")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Team AnaMALY  ·  AESS Hackathon Concept Document")
    sr.italic = True
    sr.font.size = Pt(11)
    sr.font.name = "Calibri"

    doc.add_paragraph()  # spacer

    # ===========================================================
    # 1. Project Summary
    # ===========================================================
    add_heading(doc, "1. Project Summary", level=1)
    add_table(
        doc,
        header=["Field", "Project details"],
        rows=[
            ["Project title",
             "AI-Driven Atmospheric Drag Sail for Rapid CubeSat Deorbiting"],
            ["Team name", "AnaMALY"],
            ["Problem area",
             "LEO congestion. Satellites at 600 km take decades to deorbit "
             "naturally, violating the new ESA “Zero Debris” 5-year "
             "mandate (ESSB-ST-U-007 Issue 1, Req. SDM-50)."],
            ["Proposed solution",
             "A 6U CubeSat with an integrated thin-film drag sail. The "
             "De-Orbiting Subsystem houses a Subsystem MCU (Manager) and a "
             "dedicated AI Inference IC. The IC runs a distilled deployment "
             "policy trained on the ground from real NOAA F10.7 data, "
             "selecting the optimal sail-deployment moment to maximise the "
             "drag-time integral."],
            ["Core proof",
             "Orbital decay simulation comparing No Sail (~37 yr), Naive sail "
             "(4.77 yr post-EOM), and AI-Adaptive sail (4.77 yr post-EOM "
             "with 5.30 yr 95th-percentile worst case)."],
            ["Tools used",
             "Python (custom drag integrator + scikit-learn Gaussian Process "
             "forecaster + m2cgen for embedded-C distillation), Verilog "
             "(burn-wire state machine), NOAA SWPC F10.7 archive, Three.js "
             "interactive dashboard."],
        ],
        col_widths_cm=[3.5, 12.0],
    )

    # ===========================================================
    # 2. Problem Statement
    # ===========================================================
    add_heading(doc, "2. Problem Statement", level=1)
    add_para(
        doc,
        "At 600 km the atmosphere is rarefied yet present, and conventional "
        "passive deorbiting can take 25 to 50 years — a major contributor "
        "to Kessler Syndrome risk. ESA now mandates Zero Debris reentry within "
        "5 years of mission End-of-Life. Passive sails frequently miss this "
        "target because atmospheric density at this altitude is highly "
        "volatile, tied directly to the 11-year solar cycle and to short-term "
        "geomagnetic storm events. A sail deployed at the wrong moment may "
        "deliver only a fraction of its potential drag.",
    )

    # ===========================================================
    # 3. Proposed Solution
    # ===========================================================
    add_heading(doc, "3. Proposed Solution", level=1)
    add_para(
        doc,
        "The Smart Deorbit Subsystem pairs a lightweight thin-film drag sail "
        "with an AI Inference IC governed by a Subsystem MCU (Manager) with "
        "Error Correction Code memory. Solar flux (F10.7) data is uplinked "
        "via the spacecraft’s COM block (UHF/VHF), processed on the "
        "ground into a probabilistic forecast and a distilled deployment "
        "policy, and stored on the IC pre-launch. In flight, the IC polls "
        "current state each orbit and issues WAIT or DEPLOY to the MCU; the "
        "MCU coordinates with ADCS to keep the sail perpendicular to the "
        "velocity vector during reentry, maximising drag.",
    )

    # ===========================================================
    # 4. System Architecture
    # ===========================================================
    add_heading(doc, "4. System Architecture", level=1)
    add_para(
        doc,
        "The De-Orbiting Subsystem is electrically and logically isolated "
        "from the main spacecraft bus. It receives power from the ESP and "
        "communicates with the OBC over a single CAN link, minimising fault "
        "propagation. Block names below match the published subsystem block "
        "diagram exactly.",
    )
    add_table(
        doc,
        header=["Block (per diagram)", "Function", "Implementation"],
        rows=[
            ["ESP",
             "Power conditioning (3.3 V / 5 V) from batteries and solar panels.",
             "Spacecraft electrical power subsystem"],
            ["OBC",
             "Main flight computer; CAN link to deorbit subsystem.",
             "Bus-class on-board computer"],
            ["COM",
             "UHF downlink / VHF uplink, including ground-supplied F10.7 forecasts.",
             "Analog UHF + VHF radio chain"],
            ["ADCS",
             "Attitude and position determination; Max-Drag attitude lock at reentry.",
             "Sun Sensors (analog) + Gyroscope (I²C) + Magnetometer (SPI) "
             "+ Magnetorquers (PWM) + GPS (UART)"],
            ["Subsystem MCU (Manager)",
             "Coordinates the AI Inference IC and arms the Deployment Driver; "
             "stores critical parameters in ECC memory.",
             "RISC-V manager with Error Correction Code memory"],
            ["AI Inference IC",
             "Runs the distilled deployment policy each orbit and returns "
             "WAIT/DEPLOY to the MCU.",
             "Dedicated AI Inference IC (Movidius Myriad 2 / Coral / Akida class)"],
            ["Watchdog Timer",
             "Power-On Reset and Restart on IC lock-up; long-duration EOL "
             "fail-safe.",
             "Independent watchdog with 365-day heartbeat timeout"],
            ["Deployment Driver",
             "High-current pulse output to heat the burn wires.",
             "Electrically isolated MOSFET driver with current limiting"],
            ["Deployment Mechanism",
             "Physical release of the stowed sail.",
             "Dual burn-wire / Nichrome (redundant)"],
        ],
        col_widths_cm=[4.5, 6.0, 5.5],
    )

    doc.add_paragraph()
    add_figure(doc, BLOCK_DIAGRAM,
        "Figure 1 — De-Orbiting Subsystem block diagram. The dashed box "
        "denotes the subsystem boundary; only ESP power and a single CAN "
        "link to the OBC cross it. Watchdog Timer asserts POR on the "
        "Subsystem MCU (Manager); the AI Inference IC reports WAIT/DEPLOY "
        "to the MCU, which arms the Deployment Driver to fire the dual "
        "burn-wire deployment mechanism.",
        width_cm=15.0)

    # ===========================================================
    # 5. Operating Logic (Adaptive Algorithm)
    # ===========================================================
    add_heading(doc, "5. Operating Logic (Adaptive Algorithm)", level=1)
    add_numbered(doc, [
        "Mission Phase: the satellite performs its primary mission while the "
        "AI Inference IC is power-gated except for periodic health checks.",

        "EOL Trigger: the mission ends, or the Watchdog Timer expires from "
        "loss of ground heartbeat (Dead-man Switch).",

        "Data Harvest: the OBC requests the latest NOAA F10.7 outlook via "
        "COM and forwards orbital state from ADCS (GPS) to the Subsystem "
        "MCU over CAN.",

        "AI Inference (two-stage). Ground segment: a Gaussian Process "
        "forecaster, trained on 21 years of NOAA SWPC F10.7 observations, "
        "produces a probabilistic forecast (median + 5–95 % uncertainty "
        "band). A robust optimizer evaluates each candidate deployment day "
        "across 30 sampled futures and selects the day minimising expected "
        "post-EOM lifetime. Flight segment: that policy is distilled into a "
        "depth-5 decision tree, exported to 970 bytes of portable C, and "
        "executed on the AI Inference IC each orbit. Inputs (5): mission "
        "day, current F10.7, F10.7 30-day slope, altitude, days since EOM. "
        "Output: WAIT or DEPLOY. Inference time: microseconds on a "
        "Cortex-M-class part.",

        "Sequential Verification: the Subsystem MCU requires two independent "
        "DEPLOY signals from the AI Inference IC, spaced 60 s apart, before "
        "arming the Deployment Driver. This protocol rejects single-event "
        "transients and bit-flip false triggers.",

        "Active Aerobraking: ADCS locks the sail in a Max-Drag attitude "
        "(perpendicular to the velocity vector) and the Subsystem MCU "
        "monitors decay until reentry.",
    ])

    # ===========================================================
    # 6. Drag and Decay Model
    # ===========================================================
    add_heading(doc, "6. Drag and Decay Model", level=1)
    add_para(doc,
        "The simulation integrates the secular orbital-decay equation under "
        "an exponential atmosphere with F10.7-modulated density. The "
        "instantaneous drag force per pass is:",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("F_x  =  ½ · ρ(h, F10.7) · v² · A_eff · C_d")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.italic = True

    add_para(doc,
        "and the secular semi-major-axis decay (circular-orbit assumption) is:",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("da/dt  =  − (C_d · A_eff / m) · ρ · v · a")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.italic = True

    add_para(doc,
        "where v = √(μ/a), a = R_⊕ + h, and ρ(h, F10.7) "
        "is the F10.7-modulated atmospheric density. The integrator uses an "
        "adaptive sub-step that caps altitude change per step at 2 km, "
        "preserving stability through the rapid late-stage decay below "
        "200 km. Drag coefficient C_d = 2.2 (free-molecular flat-plate, "
        "Moe & Moe 2005).",
    )

    # ===========================================================
    # 7. Deorbit KPI Comparison (Simulation Output)
    # ===========================================================
    add_heading(doc, "7. Deorbit KPI Comparison (Simulation Output)", level=1)
    add_para(doc,
        "Simulated outcomes for a 6U CubeSat at 600 km circular SSO, mass "
        "8 kg, 0.40 m² deployed sail (≈13× stowed area). The "
        "p95 column is the 95th-percentile worst-case post-EOM lifetime "
        "across 30 GP forecast samples; this bound is unique to the "
        "AI-driven approach because it requires the probabilistic "
        "forecast.",
    )
    add_table(
        doc,
        header=["Scenario", "Strategy",
                "Post-EOM (yr)", "p95 worst-case (yr)", "ESA 5-yr"],
        rows=[
            ["Baseline", "No sail (natural decay)",
             "≈37", "—", "FAIL"],
            ["Naive", "Sail deploys at EOM, no timing",
             "4.77", "5.30", "PASS"],
            ["AI-Adaptive (point)", "Best day under median forecast",
             "4.78", "5.30", "PASS"],
            ["AI-Adaptive (robust)", "Min E[score] across 30 GP samples",
             "4.77", "5.30", "PASS"],
            ["AI-Adaptive (risk-averse)", "Min p95 worst-case",
             "4.77", "5.30", "PASS"],
        ],
        col_widths_cm=[4.0, 5.0, 2.5, 2.5, 2.0],
    )
    add_para(doc,
        "The AI-Adaptive cases achieve the same expected post-EOM lifetime "
        "as Naive while providing a quantified worst-case bound under solar "
        "forecast uncertainty. The crowded-LEO band (400–500 km) "
        "residence drops from ≈5.8 satellite-years (no sail) to "
        "≈0.4 satellite-years (any sail strategy) — a 15× "
        "reduction in conjunction-risk contribution to that orbital shell.",
        italic=True,
    )

    # ===========================================================
    # 8. Simulation Validation
    # ===========================================================
    add_heading(doc, "8. Simulation Validation", level=1)
    add_para(doc,
        "Four figures generated by the Python simulator validate each stage "
        "of the pipeline:",
    )

    add_figure(doc, FIG_ALT,
        "Figure 2 — Altitude vs time for all five scenarios. Baseline "
        "(grey, no sail) decays slowly over decades; the four sail "
        "strategies overlap and reach reentry near year 8 from launch. "
        "The pink band marks the 400–500 km crowded LEO shell.",
        width_cm=15.0)

    add_figure(doc, FIG_SOLAR,
        "Figure 3 — Trained Gaussian Process forecaster. Black dots are "
        "21 years of real NOAA SWPC monthly F10.7 observations (Oct 2004 – "
        "Apr 2026, covering solar cycle 24 and the rising/peak phase of "
        "cycle 25). The blue line and shaded band are the GP median and "
        "5–95 % credible interval over the next 12 years; the band fans "
        "out at longer horizons, reflecting calibrated forecast uncertainty.",
        width_cm=15.0)

    add_figure(doc, FIG_KPI,
        "Figure 4 — KPI comparison. Left: total mission lifetime per "
        "scenario, with the ESA 25-yr and 5-yr post-EOM limits annotated. "
        "Right: years spent in the 400–500 km crowded LEO band, "
        "showing the ~15× reduction in conjunction-risk contribution "
        "delivered by any sail strategy.",
        width_cm=15.0)

    add_figure(doc, FIG_SCORE,
        "Figure 5 — Optimizer score landscape across the ±6-month "
        "flexibility window. Green: expected score across 30 GP forecast "
        "samples (used by the robust strategy, with ±1σ shaded band). "
        "Purple: score under the median forecast (used by the point "
        "strategy). Cyan: 95th-percentile worst-case score (used by the "
        "risk-averse strategy). The dashed vertical line marks the "
        "selected deployment day.",
        width_cm=15.0)

    # ===========================================================
    # 9. Implementation Evidence
    # ===========================================================
    add_heading(doc, "9. Implementation Evidence", level=1)
    add_bullets(doc, [
        "Python simulation of altitude drop-off using a Vallado piecewise-"
        "exponential atmosphere with F10.7 modulation calibrated to "
        "NRLMSIS-2.1 published values. Output: 5 deorbit trajectories with "
        "shaded uncertainty bands across the solar cycle.",

        "Subsystem block diagram showing the De-Orbiting Subsystem inside "
        "its own dashed boundary, with ESP-supplied 3.3 V / 5 V power and "
        "a single CAN link to the OBC — limited interfaces minimise "
        "fault propagation from the main bus.",

        "Pseudocode for the dual-confirmation trigger, demonstrating "
        "rejection of single-event transients via two independent DEPLOY "
        "signals separated by 60 seconds.",

        "Verilog module describing the Deployment Driver burn-wire "
        "state machine and its timing relationships, including current-"
        "limit protection during heating.",

        "Gaussian Process forecaster (scikit-learn; periodic + RBF + white-"
        "noise kernel) trained on 21 years of NOAA SWPC monthly F10.7 — "
        "produces probabilistic forecasts feeding the AI Inference IC policy.",

        "Distilled flight policy: 970-byte portable C decision tree generated "
        "from the optimizer via the m2cgen library. Cross-compilable for "
        "the AI Inference IC; 5-input, depth-5, deterministic, fully "
        "auditable.",

        "Interactive web dashboard with 3D orbit visualization (Three.js), "
        "live KPI table, and onboard policy display — see "
        "web/index.html.",
    ])

    # ===========================================================
    # 10. Reliability and Fail-Safes
    # ===========================================================
    add_heading(doc, "10. Reliability and Fail-Safes", level=1)
    add_bullets(doc, [
        "Radiation Resilience: the Subsystem MCU stores critical parameters "
        "in ECC (Error Correction Code) memory. The Two-YES protocol "
        "guarantees that a single bit-flip in the AI inference output "
        "cannot trigger the sail.",

        "Watchdog — Dual Function: the Watchdog Timer handles both "
        "routine MCU restart on hang and a long-duration EOL fail-safe. If "
        "no ground heartbeat is received for 365 days, it forces deployment "
        "automatically, bypassing AI logic.",

        "Mechanical Redundancy: dual burn-wires guarantee deployment even "
        "if one primary heating element fails.",

        "Subsystem Isolation: the De-Orbiting Subsystem lives inside its "
        "own boundary, drawing power separately from the ESP and "
        "communicating over a single CAN link, so a fault on the main bus "
        "cannot block or accidentally trigger sail release.",

        "AI Auditability: the policy on the AI Inference IC is a "
        "deterministic 970-byte decision tree, not a black-box neural "
        "network. Every deployment decision can be reproduced and reviewed "
        "pre-flight; the 5 input features and 4 leaf outcomes are "
        "documented in the auto-generated C source.",
    ])

    # ===========================================================
    # 11. Impact Statement
    # ===========================================================
    add_heading(doc, "11. Impact Statement", level=1)
    add_para(doc,
        "By reducing orbital graveyard time by over 80 % versus the no-sail "
        "baseline and providing a quantified worst-case bound under solar "
        "forecast uncertainty, this system aligns directly with the ESA Zero "
        "Debris Charter 2030. It transforms deorbiting from a passive wait-"
        "and-hope process into a managed, high-reliability maneuver, "
        "keeping the 600 km orbital shell sustainable for future missions "
        "and reducing collision-cascade risk for every operator that shares "
        "the band.",
    )
    add_para(doc,
        "The strongest part of this project is not the complexity of the AI. "
        "It is the clarity of the trigger logic, the transparency of the "
        "simulation assumptions, and the layered fail-safes — ECC "
        "memory, dual confirmation, redundant burn-wires, subsystem "
        "isolation, dual-function watchdog, and a fully auditable distilled "
        "policy — that together make the design credible as a flight "
        "subsystem rather than a paper concept.",
    )

    # ===========================================================
    # 12. References
    # ===========================================================
    add_heading(doc, "12. References", level=1)
    add_bullets(doc, [
        "ESA ESSB-ST-U-007 Issue 1 — Space Debris Mitigation "
        "Requirements (Req. SDM-50, in force 2023-11-01). "
        "https://technology.esa.int/upload/media/ESSB-ST-U-007-Issue1.pdf",

        "FCC 22-74 — 5-Year Deorbit Rule (September 2022). "
        "https://docs.fcc.gov/public/attachments/FCC-22-74A1.pdf",

        "NOAA SWPC — Observed Solar Cycle Indices (monthly F10.7). "
        "https://services.swpc.noaa.gov/json/solar-cycle/observed-solar-cycle-indices.json",

        "Underwood, C. et al. (2019). InflateSail de-orbit flight "
        "demonstration results. Acta Astronautica 162, 257–267. "
        "https://doi.org/10.1016/j.actaastro.2019.05.054",

        "Bonin, G., Risi, B., Zee, R. (2018). CanX-7 Drag Sail Deorbit "
        "Mission. AIAA/USU Conference on Small Satellites SSC18-X-03.",

        "Furano, G. et al. (2020). Towards the Use of Artificial "
        "Intelligence on the Edge in Space Systems. IEEE A&E Systems "
        "Magazine. https://doi.org/10.1109/MAES.2020.3008468",

        "Vallado, D. (2013). Fundamentals of Astrodynamics and Applications, "
        "4th ed. Microcosm Press. (Atmosphere model: Table 8-4.)",

        "Moe, K. & Moe, M. M. (2005). Gas–surface interactions and "
        "satellite drag coefficients. Planetary and Space Science 53, "
        "793–801.",

        "ESA Phi-Sat-1 — Intel Movidius Myriad 2, first AI in orbit "
        "(2020). https://www.esa.int/Applications/Observing_the_Earth/Ph-sat",

        "BayesWatch m2cgen — sklearn to portable C codegen. "
        "https://github.com/BayesWitnesses/m2cgen",
    ])

    # ---- Key Assumptions box ----
    doc.add_paragraph()  # spacer
    add_assumptions_box(doc,
        "Solar-cycle averaging uses the trained Gaussian Process forecaster "
        "fitted to NOAA F10.7 monthly observations from October 2004 to April "
        "2026 (covering solar cycle 24 and the rising/peak phase of cycle "
        "25). Sail deployment is modelled as instantaneous and the sail is "
        "treated as a flat-plate at C_d = 2.2 with full attitude lock. "
        "Ballistic coefficient is held constant post-deployment. The "
        "atmosphere model is a Vallado piecewise-exponential profile with "
        "an empirical F10.7 multiplier producing ~10× density swing at "
        "600 km between solar minimum (F10.7 ≈ 70 SFU) and maximum "
        "(F10.7 ≈ 200 SFU), consistent with NRLMSIS-2.1.",
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"  size: {os.path.getsize(OUT)/1024:.1f} KB")


if __name__ == "__main__":
    main()
